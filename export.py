import os
import json
from datetime import datetime
from docx import Document
import markdown

class DialogueExporter:
    @staticmethod
    def to_json(dialogue_history, filepath):
        """导出对话历史为JSON格式"""
        with open(filepath, 'w', encoding='utf-8') as f:
            json.dump(dialogue_history, f, ensure_ascii=False, indent=2)
            
    @staticmethod
    def to_markdown(dialogue_history, filepath):
        """导出对话历史为Markdown格式"""
        with open(filepath, 'w', encoding='utf-8') as f:
            f.write("# 对话记录\n\n")
            f.write(f"导出时间：{datetime.now().strftime('%Y-%m-%d %H:%M:%S')}\n\n")
            
            for msg in dialogue_history:
                f.write(f"## {msg['role']}\n")
                f.write(f"{msg['content']}\n\n")
                f.write(f"*时间：{msg['timestamp']}*\n\n")
                
    @staticmethod
    def to_docx(dialogue_history, filepath):
        """导出对话历史为Word文档格式"""
        doc = Document()
        doc.add_heading('对话记录', 0)
        
        doc.add_paragraph(f"导出时间：{datetime.now().strftime('%Y-%m-%d %H:%M:%S')}")
        doc.add_paragraph()
        
        for msg in dialogue_history:
            doc.add_heading(msg['role'], level=2)
            doc.add_paragraph(msg['content'])
            doc.add_paragraph(f"时间：{msg['timestamp']}", style='Subtle Emphasis')
            doc.add_paragraph()
            
        doc.save(filepath)
        
    @staticmethod
    def to_html(dialogue_history, filepath):
        """导出对话历史为HTML格式"""
        md_content = "# 对话记录\n\n"
        md_content += f"导出时间：{datetime.now().strftime('%Y-%m-%d %H:%M:%S')}\n\n"
        
        for msg in dialogue_history:
            md_content += f"## {msg['role']}\n"
            md_content += f"{msg['content']}\n\n"
            md_content += f"*时间：{msg['timestamp']}*\n\n"
            
        html_content = markdown.markdown(md_content)
        
        with open(filepath, 'w', encoding='utf-8') as f:
            f.write("""
<!DOCTYPE html>
<html>
<head>
    <meta charset="utf-8">
    <title>对话记录</title>
    <style>
        body { font-family: Arial, sans-serif; max-width: 800px; margin: 0 auto; padding: 20px; }
        h1 { color: #2c3e50; }
        h2 { color: #34495e; }
        .timestamp { color: #7f8c8d; font-style: italic; }
    </style>
</head>
<body>
""")
            f.write(html_content)
            f.write("\n</body>\n</html>")
